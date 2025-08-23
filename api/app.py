# backend/app.py - Real Agentic RAG Implementation

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import json
import asyncio
import os
import subprocess
import tiktoken
from dotenv import load_dotenv
import uuid
import io
import time
import re
import hashlib
import requests

# Load environment variables
load_dotenv()

# Import our existing components from the notebook
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain_community.document_loaders import DirectoryLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from operator import itemgetter
from langchain_core.output_parsers import StrOutputParser
from langchain_core.tools import tool
from langchain_community.tools import TavilySearchResults
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
# LangChain's document class name changed across versions; prefer minimal shim
try:
    from langchain_core.documents import Document as LCDocument
except Exception:
    LCDocument = None
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode
from typing_extensions import TypedDict, Annotated
import arxiv

# Import file processing libraries
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image

# Global storage for uploaded lab data
uploaded_lab_data = {}

# File processing functions
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file with enhanced error handling"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        
        print(f"DEBUG: PDF has {len(pdf_reader.pages)} pages")  # Debug line
        
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            print(f"DEBUG: Page {i+1} extracted text length: {len(page_text)}")  # Debug line
            if page_text:
                text += page_text + "\n"
        
        if not text.strip():
            print("DEBUG: No text extracted from PDF, trying alternative method")  # Debug line
            # Try to get any metadata or info
            try:
                info = pdf_reader.metadata
                text = f"PDF Metadata: {info}\nNo readable text content found. This might be a scanned PDF or image-based document."
            except:
                text = "PDF file uploaded but no readable text content could be extracted. This might be a scanned document or image-based PDF."
        
        return text
    except Exception as e:
        print(f"DEBUG: PDF extraction error: {str(e)}")  # Debug line
        return f"Error processing PDF: {str(e)}. Please ensure the PDF contains readable text."

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")

def extract_text_from_image(file_content: bytes) -> str:
    """Extract text from image using OCR or return image info"""
    try:
        image = Image.open(io.BytesIO(file_content))
        
        # Try OCR first (if pytesseract is available)
        try:
            import pytesseract
            text = pytesseract.image_to_string(image)
            if text.strip():
                return text
        except (ImportError, Exception):
            pass  # OCR failed or not available, continue with image info
        
        # If OCR fails or no text found, return image information
        width, height = image.size
        format_info = image.format or "Unknown"
        mode_info = image.mode
        
        return f"Image file: {format_info} format, {width}x{height} pixels, {mode_info} color mode. No text extracted. Please describe the lab results in your question."
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing image: {str(e)}")

def process_uploaded_file(file: UploadFile, max_mb: int = 10) -> str:
    """Process uploaded file and extract text"""
    # Validate content type and size
    allowed_content_types = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
    }

    if file.content_type not in allowed_content_types:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF or DOCX.")

    content = file.file.read()
    file.file.seek(0)  # Reset file pointer
    max_bytes = max_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(status_code=400, detail=f"File too large. Maximum allowed size is {max_mb} MB.")
    
    try:
        if file.filename.lower().endswith('.pdf'):
            return extract_text_from_pdf(content)
        elif file.filename.lower().endswith(('.docx', '.doc')):
            return extract_text_from_docx(content)
            # JPG/PNG support removed - only text-based files supported
    # elif file.filename.lower().endswith(('.jpg', '.jpeg', '.png')):
    #     return extract_text_from_image(content)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
    except Exception as e:
        # Return a helpful message instead of failing completely
        return f"File uploaded but content could not be extracted: {str(e)}. Please describe the lab results in your question."

# Initialize FastAPI app
app = FastAPI(title="Personal Health Copilot API", version="1.0.0")

# CORS middleware for frontend communication (env-driven, defaults preserve localhost behavior)
default_origins = ["http://localhost:3000", "http://localhost:3001"]
env_origins = os.getenv("ALLOWED_ORIGINS")
allow_origins = (
    [origin.strip() for origin in env_origins.split(",") if origin.strip()]
    if env_origins
    else default_origins
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------------
# Optional privacy/compliance toggles (default safe/off)
# -----------------------------
GDPR_MODE = os.getenv("GDPR_MODE", "false").lower() in ("1", "true", "yes")
COMPLIANCE_MODE = os.getenv("COMPLIANCE_MODE", "false").lower() in ("1", "true", "yes")

# -----------------------------
# Trusted sources mode (allow-list + freshness)
# -----------------------------
TRUSTED_SOURCES_MODE = os.getenv("TRUSTED_SOURCES_MODE", "false").lower() in ("1", "true", "yes")
_default_domains = [
    "mayoclinic.org",
    "medlineplus.gov",
    "cdc.gov",
    "nih.gov",
    "fda.gov",
    "who.int",
]
TRUSTED_ALLOWED_DOMAINS = [
    d.strip().lower() for d in os.getenv("TRUSTED_ALLOWED_DOMAINS", ",".join(_default_domains)).split(",") if d.strip()
]
TRUSTED_MAX_YEARS = int(os.getenv("TRUSTED_MAX_YEARS", "5"))

def _extract_domain(url: str) -> str:
    try:
        from urllib.parse import urlparse
        netloc = urlparse(url).netloc.lower()
        return netloc.split(":")[0]
    except Exception:
        return ""

def _is_domain_allowed(url: str) -> bool:
    if not TRUSTED_SOURCES_MODE:
        return True
    domain = _extract_domain(url)
    return any(domain.endswith(allowed) for allowed in TRUSTED_ALLOWED_DOMAINS)

def _parse_markdown_links(markdown_text: str) -> List[Dict[str, str]]:
    links = []
    try:
        pattern = re.compile(r"\[([^\]]+)\]\(([^\)]+)\)")
        for match in pattern.finditer(markdown_text or ""):
            text = match.group(1)
            url = match.group(2)
            links.append({"text": text, "url": url, "domain": _extract_domain(url)})
    except Exception:
        pass
    return links

def _redact_pii(text: str) -> str:
    """Lightweight PII redaction (emails, phones, simple IDs). Non-destructive; demo-safe.
    This is not a substitute for production-grade de-identification.
    """
    if not isinstance(text, str) or not text:
        return text
    result = text
    # Emails
    result = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "[redacted-email]", result)
    # Phone numbers (international and common patterns)
    result = re.sub(r"(?:(?:\+?\d{1,3}[\s.-]?)?(?:\(\d{2,4}\)[\s.-]?)?\d{3}[\s.-]?\d{3,4}[\s.-]?\d{3,4})", "[redacted-phone]", result)
    # Simple MRN-like IDs (e.g., AA-1234 or ABC-0001)
    result = re.sub(r"\b[A-Z]{2,4}-\d{3,6}\b", "[redacted-id]", result)
    # Long digit sequences (credit/ID-like)
    result = re.sub(r"\b\d{10,}\b", "[redacted-number]", result)
    return result

def _audit_log(action: str, request: Request = None, extra: Optional[Dict[str, Any]] = None) -> None:
    """Minimal audit logging to stdout when compliance is enabled."""
    if not (GDPR_MODE or COMPLIANCE_MODE):
        return
    try:
        client_ip = request.client.host if request and request.client else "unknown"
        entry = {
            "action": action,
            "client_ip": client_ip,
            "path": getattr(request, "url", None).path if request else None,
            "method": getattr(request, "method", None),
            "ts": time.time(),
        }
        if extra:
            entry.update(extra)
        print(f"AUDIT {json.dumps(entry)}")
    except Exception:
        # Never fail request due to audit logging
        pass

# Pydantic models
class ChatMessage(BaseModel):
    message: str
    user_id: str = "default"
    conversation_history: List[Dict[str, str]] = []  # Add conversation history
    session_id: Optional[str] = None  # Add session_id for lab results

class GenerateTitleRequest(BaseModel):
    messages: List[Dict[str, str]]  # List of messages with sender and text

class AgentActivity(BaseModel):
    tool_used: str
    status: str
    details: Dict[str, Any]

class UploadResponse(BaseModel):
    session_id: str
    files_processed: int
    message: str

class LabResultsRequest(BaseModel):
    session_id: str

# New Pydantic models for additive enterprise endpoints
class ConversationMessage(BaseModel):
    sender: str
    text: str

class PatientSummaryRequest(BaseModel):
    conversation_history: List[ConversationMessage] = []
    session_id: Optional[str] = None

class PatientSummaryResponse(BaseModel):
    problems: List[str] = []
    medications: List[str] = []
    labs: Dict[str, Any] = {}
    plan: List[str] = []

class LabsInterpretRequest(BaseModel):
    session_id: str
    question: Optional[str] = None

class LabsInterpretResponse(BaseModel):
    flags: List[str] = []
    explanations: List[str] = []
    recommendations: List[str] = []

class ProviderNoteRequest(BaseModel):
    messages: List[ConversationMessage]
    sections: Optional[List[str]] = None  # e.g., ["HPI", "Assessment", "Plan"]
    snapshot: Optional[Dict[str, Any]] = None  # demo patient snapshot (problems/medications/last_labs)
    notes: Optional[str] = None  # visit notes from provider UI

class ProviderNoteResponse(BaseModel):
    note: Dict[str, str]
    markdown: str

# Global variables for medical system
split_docs = None
bm25_retriever = None
rag_chain = None
compiled_health_graph = None

# Initialize LLM
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.1,
    openai_api_key=os.getenv("OPENAI_API_KEY")
)

# -----------------------------
# Simple in-memory rate limiting (guest mode)
# -----------------------------
CHAT_LIMIT_PER_MINUTE = int(os.getenv("CHAT_LIMIT_PER_MINUTE", "20"))
UPLOAD_LIMIT_PER_MINUTE = int(os.getenv("UPLOAD_LIMIT_PER_MINUTE", "5"))

rate_counters: Dict[str, list] = {}

def _now() -> float:
    return time.time()

def _clean_old_requests(timestamps: list, window: float) -> None:
    cutoff = _now() - window
    while timestamps and timestamps[0] < cutoff:
        timestamps.pop(0)

def check_rate_limit(key: str, limit: int, window_seconds: float = 60.0) -> bool:
    """Tokenless sliding window limiter stored in memory. Returns True if allowed."""
    timestamps = rate_counters.setdefault(key, [])
    _clean_old_requests(timestamps, window_seconds)
    if len(timestamps) >= limit:
        return False
    timestamps.append(_now())
    return True

# -----------------------------
# Simple in-memory TTL cache (Vercel/serverless-friendly per-instance cache)
# -----------------------------
CACHE_TTL_SECONDS = int(os.getenv("CACHE_TTL_SECONDS", "300"))  # 5 minutes default
_cache_store: Dict[str, Dict[str, Any]] = {}

def _cache_get(key: str) -> Optional[Any]:
    if not CACHE_TTL_SECONDS:
        return None
    entry = _cache_store.get(key)
    if not entry:
        return None
    if _now() - entry.get("ts", 0) > CACHE_TTL_SECONDS:
        try:
            del _cache_store[key]
        except Exception:
            pass
        return None
    return entry.get("value")

def _cache_set(key: str, value: Any) -> None:
    if not CACHE_TTL_SECONDS:
        return
    _cache_store[key] = {"ts": _now(), "value": value}

def _hash_key(payload: Any) -> str:
    try:
        blob = json.dumps(payload, sort_keys=True, ensure_ascii=False)
    except Exception:
        blob = str(payload)
    return hashlib.sha256(blob.encode("utf-8")).hexdigest()

# Helper: classify triage level for safety signaling
def classify_triage_level(user_message: str, lab_context: str = "") -> Dict[str, Any]:
    """Classify into routine | urgent | emergency with brief reasons. Returns dict.
    Uses the existing compiled_health_graph with a constrained JSON instruction.
    Fallback: heuristic classification.
    """
    try:
        instruction = (
            "Classify medical urgency for the user's concern. Return STRICT JSON with keys: "
            "level (one of 'routine','urgent','emergency') and reasons (array of 1-3 short strings). "
            "Consider any lab context if provided. Do not include any extra text."
        )
        lab_section = f"\nLAB CONTEXT:\n{lab_context}" if lab_context else ""
        prompt = f"{instruction}{lab_section}\n\nQUESTION: {user_message}"
        system_message = SystemMessage(content="You are a clinical triage assistant. Output JSON only.")
        triage_resp = compiled_health_graph.invoke({
            "messages": [system_message, HumanMessage(content=prompt)]
        })
        content = triage_resp["messages"][-1].content
        try:
            parsed = json.loads(content)
        except Exception:
            start = content.find("{")
            end = content.rfind("}")
            parsed = json.loads(content[start:end+1]) if start != -1 and end != -1 else {}
        level = parsed.get("level", "routine").lower()
        if level not in ["routine", "urgent", "emergency"]:
            level = "routine"
        reasons = parsed.get("reasons", [])
        if not isinstance(reasons, list):
            reasons = [str(reasons)]
        return {"level": level, "reasons": reasons[:3]}
    except Exception:
        # Heuristic fallback
        text = (user_message or "").lower()
        emergency_terms = [
            "chest pain", "shortness of breath", "severe bleeding", "stroke", "fainting",
            "confusion", "not breathing", "blue lips", "unconscious", "suicidal"
        ]
        urgent_terms = [
            "high fever", "severe pain", "worsening", "can't keep fluids", "dehydration",
            "new rash", "persistent vomiting", "pregnancy complications"
        ]
        if any(t in text for t in emergency_terms):
            return {"level": "emergency", "reasons": ["Potential emergency symptom detected."]}
        if any(t in text for t in urgent_terms):
            return {"level": "urgent", "reasons": ["Symptoms may require timely clinical care."]}
        return {"level": "routine", "reasons": ["General guidance; monitor and follow up."]}

# RAG prompt template (from Task 4)
RAG_TEMPLATE = """\
You are a helpful and knowledgeable medical assistant. Use the context provided below to answer the medical question. If you don't know the answer, say so. Don't make up information.

IMPORTANT: When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href="...">. Keep responses clean and readable.

### Question
{question}

### Context
{context}

### Answer
"""

rag_prompt = ChatPromptTemplate.from_template(RAG_TEMPLATE)

# Medical tools (from Task 4)
@tool
def retrieve_medical_information(
    query: str
):
    """Use RAG to retrieve medical information from verified sources"""
    if bm25_retriever:
        docs = bm25_retriever.get_relevant_documents(query)
        return "\n".join([doc.page_content for doc in docs])
    return "Medical information retrieved from verified sources"

@tool
def web_search_medical_info(
    query: str
):
    """Search the web for current medical information, treatments, and health research"""
    search_tool = TavilySearchResults(max_results=5)
    results = search_tool.invoke(query)
    # Filter by allow-list when trusted mode is on
    usable = [r for r in results if _is_domain_allowed(r.get("url", ""))] if TRUSTED_SOURCES_MODE else results
    if not usable:
        return (
            f"No trusted results found for '{query}'."
            if TRUSTED_SOURCES_MODE
            else f"No results found for '{query}'."
        )
    formatted_results = []
    for result in usable[:3]:
        formatted_results.append(
            f"Title: {result.get('title','')}\nURL: {result.get('url','')}\nContent: {result.get('content','')}\n"
        )
    return f"Web search results for '{query}':\n" + "\n".join(formatted_results)

@tool
def search_medical_research(
    query: str
):
    """Search ArXiv for medical research papers and scientific studies"""
    search = arxiv.Search(
        query=query,
        max_results=3,
        sort_by=arxiv.SortCriterion.Relevance
    )
    results = []
    for result in search.results():
        label = " [Preprint]" if TRUSTED_SOURCES_MODE else ""
        results.append(f"Title: {result.title}{label}\nAuthors: {', '.join([author.name for author in result.authors])}\nAbstract: {result.summary[:300]}...\nURL: {result.entry_id}\n")
    return f"Medical research results for '{query}':\n" + "\n".join(results)

@tool
def search_pubmed(query: str):
    """Search PubMed (peer-reviewed). Returns up to 3 citations with title and URL."""
    try:
        params = {"db": "pubmed", "term": query, "retmax": 3}
        es = requests.get(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi",
            params={**params, "retmode": "json"},
            timeout=10,
        )
        ids = (es.json().get("esearchresult", {}).get("idlist", []) if es.ok else [])
        if not ids:
            return "No PubMed results."
        summ = requests.get(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi",
            params={"db": "pubmed", "id": ",".join(ids), "retmode": "json"},
            timeout=10,
        )
        items = (summ.json().get("result", {}) if summ.ok else {})
        formatted = []
        for pid in ids:
            m = items.get(pid, {})
            title = m.get("title", "")
            year = (m.get("pubdate", "").split(" ") or [""])[0]
            url = f"https://pubmed.ncbi.nlm.nih.gov/{pid}/"
            formatted.append(f"Title: {title} ({year})\nURL: {url}\n")
        return "PubMed results:\n" + "\n".join(formatted)
    except Exception as e:
        return f"PubMed error: {str(e)}"

@tool
def analyze_symptoms(
    symptoms: str
):
    """Analyze symptoms and suggest possible conditions based on medical knowledge"""
    return f"Based on symptoms: {symptoms}, I recommend consulting a healthcare provider for proper diagnosis."

@tool
def check_drug_interactions(
    medications: str
):
    """Check for potential drug interactions with current medications"""
    return f"Checking interactions for medications: {medications}. Please consult a pharmacist for complete safety information."

@tool
def interpret_lab_results(
    lab_results: str
):
    """Interpret lab test results and explain what they mean in plain language"""
    return f"Interpreting lab results: {lab_results}. Please consult with your healthcare provider for complete medical interpretation."

@tool
def analyze_uploaded_lab_results(
    lab_data: str,
    user_question: str
):
    """Analyze uploaded lab results and answer user questions about them"""
    return f"Analyzing lab results: {lab_data[:200]}... for question: {user_question}. This provides medical interpretation of uploaded lab data."

# Initialize medical documents and RAG system
def initialize_medical_system():
    """Initialize medical documents and RAG system (from Task 3)"""
    global split_docs, bm25_retriever, rag_chain, compiled_health_graph
    
    print("Initializing medical system...")
    
    use_embedded = os.getenv("USE_EMBEDDED_CORPUS", "true").lower() in ("1", "true", "yes") or os.getenv("DISABLE_STARTUP_DOWNLOADS", "").lower() in ("1", "true", "yes")

    if use_embedded:
        print("Using embedded medical corpus (serverless-friendly)...")
        embedded_texts = [
            ("Mayo Clinic — Diabetes", "Diabetes is a chronic condition affecting how your body turns food into energy. Symptoms may include increased thirst, frequent urination, fatigue. Lifestyle and medications can help manage blood sugar."),
            ("Mayo Clinic — Hypertension", "High blood pressure often has no symptoms but increases risk of heart disease and stroke. Lifestyle changes and medications are common treatments."),
            ("MedlinePlus — CBC", "A complete blood count (CBC) measures red blood cells, white blood cells, hemoglobin, and platelets. Abnormal levels can indicate infection, anemia, or other conditions."),
            ("FDA — Drug Interactions", "Drug interactions can change how medicines work or cause unexpected side effects. Always review medications with a healthcare professional."),
        ]
        # Build documents compatible with current LangChain version
        if LCDocument is not None:
            docs = [LCDocument(page_content=txt, metadata={"source": title}) for title, txt in embedded_texts]
        else:
            # Fallback: compatible structure for BM25 retriever
            docs = [type("_Doc", (), {"page_content": txt, "metadata": {"source": title}})() for title, txt in embedded_texts]
    else:
        # Create data directory
        os.makedirs("data", exist_ok=True)
        # Download medical documents from URLs (from Task 3)
        print("Downloading medical documents...")
        medical_urls = [
            "https://www.mayoclinic.org/diseases-conditions/diabetes/symptoms-causes/syc-20371444",
            "https://www.mayoclinic.org/diseases-conditions/high-blood-pressure/symptoms-causes/syc-20373410",
            "https://www.mayoclinic.org/diseases-conditions/heart-disease/symptoms-causes/syc-20353118",
            "https://www.mayoclinic.org/diseases-conditions/arthritis/symptoms-causes/syc-20350772",
            "https://www.mayoclinic.org/diseases-conditions/asthma/symptoms-causes/syc-20369653",
            "https://medlineplus.gov/druginfo/meds/a682878.html",
            "https://medlineplus.gov/druginfo/meds/a682159.html",
            "https://medlineplus.gov/druginfo/meds/a682345.html",
            "https://www.fda.gov/drugs/drug-safety-and-availability/drug-interactions",
            "https://medlineplus.gov/lab-tests/complete-blood-count-cbc/",
            "https://medlineplus.gov/lab-tests/blood-glucose-test/",
            "https://medlineplus.gov/lab-tests/cholesterol-levels/"
        ]
        for i, url in enumerate(medical_urls):
            try:
                filename = f"data/medical_doc_{i+1}.html"
                subprocess.run(["curl", "-o", filename, url], check=True)
                print(f"Downloaded: {filename}")
            except Exception as e:
                print(f"Failed to download {url}: {e}")
        # Load documents using DirectoryLoader (from Task 3)
        print("Loading medical documents...")
        path = "data/"
        loader = DirectoryLoader(path, glob="*.html")
        docs = loader.load()
    
    print(f"Loaded {len(docs)} medical documents")
    
    # Text splitter configuration (from Task 3)
    def tiktoken_len(text):
        tokens = tiktoken.encoding_for_model("gpt-4o").encode(text)
        return len(tokens)
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=200,
        length_function=tiktoken_len,
    )
    
    # Split documents
    split_docs = text_splitter.split_documents(docs)
    print(f"Split into {len(split_docs)} chunks")
    
    # Initialize BM25 retriever (from Task 6)
    print("Initializing BM25 retriever...")
    bm25_retriever = BM25Retriever.from_documents(split_docs)
    
    # Create RAG chain with BM25 (from Task 4)
    print("Creating RAG chain...")
    rag_chain = (
        {"context": itemgetter("question") | bm25_retriever, "question": itemgetter("question")}
        | rag_prompt | llm | StrOutputParser()
    )
    
    # Create agent workflow (from Task 4)
    print("Creating agent workflow...")
    
    class AgentState(TypedDict):
        messages: Annotated[list, add_messages]
    
    # Tool belt with all medical tools (PubMed added)
    tool_belt = [
        retrieve_medical_information, 
        web_search_medical_info, 
        search_medical_research, 
        analyze_symptoms, 
        check_drug_interactions, 
        interpret_lab_results,
        analyze_uploaded_lab_results,
        search_pubmed,
    ]
    
    llm_with_tools = llm.bind_tools(tool_belt)
    
    def call_model(state):
        messages = state["messages"]
        # Add system message to instruct LLM to use markdown links, and prefer trusted sources when enabled
        trust_note = " When external sources are needed, prefer PubMed/Medline and approved domains (mayoclinic.org, medlineplus.gov, cdc.gov, nih.gov, fda.gov, who.int). Label arXiv as preprint. If no trusted recent source (last 5 years) is available, say so." if TRUSTED_SOURCES_MODE else ""
        system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable." + trust_note)
        messages_with_system = [system_message] + messages
        response = llm_with_tools.invoke(messages_with_system)
        return {"messages": [response]}
    
    tool_node = ToolNode(tool_belt)
    
    health_graph = StateGraph(AgentState)
    health_graph.add_node("agent", call_model)
    health_graph.add_node("action", tool_node)
    health_graph.set_entry_point("agent")
    
    def should_continue(state):
        last_message = state["messages"][-1]
        if last_message.tool_calls:
            return "action"
        return END
    
    health_graph.add_conditional_edges("agent", should_continue)
    health_graph.add_edge("action", "agent")
    
    compiled_health_graph = health_graph.compile()
    
    print("✅ Medical system initialized successfully!")

# Initialize the system on startup
@app.on_event("startup")
async def startup_event():
    """Initialize medical system on startup"""
    initialize_medical_system()

# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.remove(websocket)

    async def send_personal_message(self, message: str, websocket: WebSocket):
        await websocket.send_text(message)

    async def broadcast(self, message: str):
        for connection in self.active_connections:
            await connection.send_text(message)

manager = ConnectionManager()

# API Routes
@app.get("/")
async def root():
    return {"message": "Personal Health Copilot API is running!"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "Personal Health Copilot API"}

@app.post("/chat")
async def chat_endpoint(chat_message: ChatMessage, request: Request):
    """Main chat endpoint for medical queries using real agent"""
    
    if not compiled_health_graph:
        return {"error": "Medical system not initialized"}
    
    # Rate limit per-IP for guest mode
    client_ip = request.client.host if request.client else "unknown"
    if not check_rate_limit(f"chat:{client_ip}", CHAT_LIMIT_PER_MINUTE):
        raise HTTPException(status_code=429, detail="Rate limit exceeded. Please wait and try again.")

    # Default category to a safe value; will be refined below.
    question_category = "HEALTH_QUESTION"

    try:
        # Optional: redact incoming message for demo privacy if toggled
        original_message = chat_message.message
        if GDPR_MODE or COMPLIANCE_MODE:
            chat_message.message = _redact_pii(chat_message.message)

        # Build conversation history from previous messages
        messages = []
        
        # Add conversation history if provided
        for msg in chat_message.conversation_history:
            if msg.get("sender") == "user":
                messages.append(HumanMessage(content=msg.get("text", "")))
            elif msg.get("sender") == "assistant":
                messages.append(AIMessage(content=msg.get("text", "")))
        
        # Add current message (already redacted if enabled)
        messages.append(HumanMessage(content=chat_message.message))
        
        # Try cache first (message + lab session + brief history)
        cache_key = _hash_key({
            "m": chat_message.message,
            "s": chat_message.session_id,
            "h": [(m.get("sender"), m.get("text")) for m in chat_message.conversation_history[-3:]],
        })
        cached = _cache_get(cache_key)
        if cached:
            _audit_log("chat_cache_hit", request, extra={"key": cache_key})
            return cached

        # Get lab context if session_id is provided
        lab_context = ""
        if chat_message.session_id and chat_message.session_id in uploaded_lab_data:
            lab_data = uploaded_lab_data[chat_message.session_id]
            print(f"DEBUG: Found lab data for session {chat_message.session_id}")  # Debug line
            lab_context = f"\n\nLAB RESULTS CONTEXT:\n"
            
            has_valid_content = False
            for file_data in lab_data["files"]:
                # Show more content for better analysis
                content_preview = file_data['content'][:2000] if len(file_data['content']) > 2000 else file_data['content']
                print(f"DEBUG: Lab content preview: {content_preview[:200]}...")  # Debug line
                
                # Check if content is meaningful (not just error messages)
                if content_preview and not content_preview.startswith("Error processing") and not content_preview.startswith("PDF file uploaded but no readable"):
                    has_valid_content = True
                    lab_context += f"File: {file_data['filename']}\nContent: {content_preview}\n\n"
                else:
                    lab_context += f"File: {file_data['filename']}\nContent: {content_preview}\n\n"
            
            if not has_valid_content:
                lab_context = ""  # Don't pass empty/invalid content to AI
                print("DEBUG: No valid content found in uploaded files")  # Debug line
        else:
            print(f"DEBUG: No lab data found for session {chat_message.session_id}")  # Debug line
        
        # Use LLM to intelligently classify the question type and generate appropriate response
        # First, let's get the actual conversation context for better classification
        conversation_context = ""
        if len(messages) > 1:
            # Get the last few messages for context
            recent_messages = messages[-3:]  # Last 3 messages for context
            conversation_context = "\n".join([
                f"{'User' if isinstance(msg, HumanMessage) else 'Assistant'}: {msg.content}"
                for msg in recent_messages[:-1]  # Exclude the current message
            ])
        
        classification_prompt = f"""
        You are a medical AI assistant. Analyze the user's message and classify it into one of these categories:

        1. GREETING - Simple greetings like "hi", "hello", "hey"
        2. HEALTH_QUESTION - Health/medical questions, symptoms, treatments, etc.
        3. NON_HEALTH_QUESTION - Non-medical topics (politics, weather, etc.)
        4. FOLLOW_UP - Requests for more details, clarification, or elaboration about ANY previous topic
        5. CONVERSATION_RESUMPTION - User wants to continue previous conversation after interruption

        CONVERSATION CONTEXT:
        {conversation_context}
        
        {lab_context}
        
        CURRENT USER MESSAGE: "{chat_message.message}"

        CLASSIFICATION RULES:
        - If user says anything indicating they want to continue, resume, or get back to previous topic → CONVERSATION_RESUMPTION
        - If user asks for more details about ANY previous topic → FOLLOW_UP  
        - If user uses transition words like "ok", "so", "anyway", "let's", "continue", "back to" → CONVERSATION_RESUMPTION
        - If user asks for clarification or elaboration → FOLLOW_UP
        - If user references previous conversation without specific health question → CONVERSATION_RESUMPTION

        Respond with ONLY: GREETING, HEALTH_QUESTION, NON_HEALTH_QUESTION, FOLLOW_UP, or CONVERSATION_RESUMPTION
        """
        
        try:
            # Use LLM to classify the question
            system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable.")
            classification_response = compiled_health_graph.invoke({
                "messages": [system_message, HumanMessage(content=classification_prompt)]
            })
            question_category = classification_response["messages"][-1].content.strip().upper()
            
            # Generate appropriate response based on classification
            if question_category == "GREETING":
                final_response_content = "Hello! How can I help you with your health today?"
            elif question_category == "CONVERSATION_RESUMPTION":
                # Continue with the previous health conversation context
                system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable." + (" Prefer PubMed/approved domains when using web sources." if TRUSTED_SOURCES_MODE else ""))
                messages_with_system = [system_message] + messages
                final_response = compiled_health_graph.invoke({"messages": messages_with_system})
                final_message = final_response["messages"][-1]
                final_response_content = final_message.content.replace("**", "").replace("*", "")
            elif question_category == "NON_HEALTH_QUESTION":
                # Extract topic intelligently using LLM
                topic_extraction_prompt = f"""
                Extract the main topic from this non-health question. Return only the topic in 2-3 words maximum.
                
                Question: "{chat_message.message}"
                
                Example: "explain about syria war" → "syria war"
                Example: "what's the weather like" → "weather"
                
                Topic:
                """
                
                system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable." + (" Prefer PubMed/approved domains when using web sources." if TRUSTED_SOURCES_MODE else ""))
                topic_response = compiled_health_graph.invoke({
                    "messages": [system_message, HumanMessage(content=topic_extraction_prompt)]
                })
                topic_text = topic_response["messages"][-1].content.strip()
                
                final_response_content = f"I'm your Personal Health Copilot, designed specifically to help with health and medical questions. I can assist you with symptoms, treatments, medications, lab results, and general health advice. Unfortunately, I cannot provide information about {topic_text} or other non-health topics, as I'm focused on medical and health-related assistance. How can I help with your health today?"
            else:
                # HEALTH_QUESTION or FOLLOW_UP - Use full medical agent workflow
                # Include lab context in the main conversation if available
                if lab_context and lab_context.strip():
                    # Create a specific instruction for lab results analysis
                    lab_instruction = f"""
                    IMPORTANT: The user has uploaded lab results. Please analyze the uploaded lab data and provide a comprehensive medical interpretation.
                    
                    LAB RESULTS DATA:
                    {lab_context}
                    
                    USER QUESTION: {chat_message.message}
                    
                    Please provide:
                    1. Medical interpretation of the lab results
                    2. What the values mean
                    3. Any concerning or normal findings
                    4. Recommendations if applicable
                    
                    Be thorough and medical in your analysis.
                    """
                    
                    # Add lab instruction as the first message
                    lab_system_message = HumanMessage(content=lab_instruction)
                    messages.insert(0, lab_system_message)
                else:
                    # No lab context - use normal medical agent for health questions
                    system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable.")
                    messages_with_system = [system_message] + messages
                    final_response = compiled_health_graph.invoke({"messages": messages_with_system})
                    final_message = final_response["messages"][-1]
                    
                    # Clean up the response for better user experience
                    cleaned_response = final_message.content
                    cleaned_response = cleaned_response.replace("**", "").replace("*", "")
                    cleaned_response = cleaned_response.replace("###", "").replace("##", "").replace("#", "")
                    
                    # Add empathetic prefix for follow-up questions using LLM
                    if question_category == "FOLLOW_UP":
                        empathy_prompt = f"""
                        Add an empathetic prefix to this medical response. Choose from:
                        - "Yes, absolutely! " (for explain/detail requests)
                        - "Of course! " (for what/how questions)
                        - "Perfect! " (for general follow-ups)
                        
                        User's follow-up: "{chat_message.message}"
                        Medical response: "{cleaned_response}"
                        
                        Return only the prefixed response.
                        """
                        
                        system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable.")
                        empathy_response = compiled_health_graph.invoke({
                            "messages": [system_message, HumanMessage(content=empathy_prompt)]
                        })
                        final_response_content = empathy_response["messages"][-1].content.strip()
                    else:
                        final_response_content = cleaned_response
                
                system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable." + (" Prefer PubMed/approved domains when using web sources." if TRUSTED_SOURCES_MODE else ""))
                messages_with_system = [system_message] + messages
                final_response = compiled_health_graph.invoke({"messages": messages_with_system})
                final_message = final_response["messages"][-1]
                
                # Clean up the response for better user experience
                cleaned_response = final_message.content
                cleaned_response = cleaned_response.replace("**", "").replace("*", "")
                cleaned_response = cleaned_response.replace("###", "").replace("##", "").replace("#", "")
                
                # Clean response (keeping original functionality)
                
                # Add empathetic prefix for follow-up questions using LLM
                if question_category == "FOLLOW_UP":
                    empathy_prompt = f"""
                    Add an empathetic prefix to this medical response. Choose from:
                    - "Yes, absolutely! " (for explain/detail requests)
                    - "Of course! " (for what/how questions)
                    - "Perfect! " (for general follow-ups)
                    
                    User's follow-up: "{chat_message.message}"
                    Medical response: "{cleaned_response}"
                    
                    Return only the prefixed response.
                    """
                    
                    empathy_response = compiled_health_graph.invoke({
                        "messages": [HumanMessage(content=empathy_prompt)]
                    })
                    final_response_content = empathy_response["messages"][-1].content.strip()
                else:
                    final_response_content = cleaned_response
                    
        except Exception as e:
            # Fallback to static logic if LLM fails
            current_message = chat_message.message.lower()
            greetings = ["hi", "hello", "hey", "good morning", "good afternoon", "good evening"]
            health_keywords = ["symptom", "pain", "ache", "fever", "headache", "medicine", "medical", "health", "treatment", "disease", "doctor"]
            
            if any(greeting in current_message for greeting in greetings):
                final_response_content = "Hello! How can I help you with your health today?"
            elif any(keyword in current_message for keyword in health_keywords) or len(messages) > 1 or any(resumption_word in current_message for resumption_word in ["ok", "so", "anyway", "let's", "continue", "back to", "resume"]):
                # Health question or follow-up
                system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable.")
                messages_with_system = [system_message] + messages
                final_response = compiled_health_graph.invoke({"messages": messages_with_system})
                final_message = final_response["messages"][-1]
                final_response_content = final_message.content.replace("**", "").replace("*", "")
                # Set fallback category appropriately
                question_category = "FOLLOW_UP" if len(messages) > 1 else "HEALTH_QUESTION"
                # Clean response (keeping original functionality)
            else:
                # Non-health question
                words = current_message.split()
                topic_words = [word for word in words if len(word) > 2 and word not in ["explain", "about", "tell", "me", "what", "is", "the"]]
                topic_text = " ".join(topic_words[:3]) if topic_words else "this topic"
                final_response_content = f"I'm your Personal Health Copilot, designed specifically to help with health and medical questions. I can assist you with symptoms, treatments, medications, lab results, and general health advice. Unfortunately, I cannot provide information about {topic_text} or other non-health topics, as I'm focused on medical and health-related assistance. How can I help with your health today?"
                question_category = "NON_HEALTH_QUESTION"
        
        # Simulate agent activity for real-time display (only for health questions)
        activities = []
        if question_category in ["HEALTH_QUESTION", "FOLLOW_UP", "CONVERSATION_RESUMPTION"]:
            activities = [
                {"tool_used": "BM25_Retriever", "status": "searching", "details": {"query": chat_message.message}},
                {"tool_used": "Medical_Agent", "status": "processing", "details": {"model": "gpt-4o-mini"}},
                {"tool_used": "Response_Generator", "status": "generating", "details": {"tokens": len(final_response_content.split())}}
            ]
        
        triage = classify_triage_level(chat_message.message, lab_context)
        # Agent-driven suggestion for showing triage banner: default true for health/follow-up/resumption
        triage_display = question_category in ["HEALTH_QUESTION", "FOLLOW_UP", "CONVERSATION_RESUMPTION"]
        # Minimal audit trail
        _audit_log("chat", request, extra={"redacted": (GDPR_MODE or COMPLIANCE_MODE), "msg_len": len(original_message or "")})
        sources = _parse_markdown_links(final_response_content)
        if TRUSTED_SOURCES_MODE:
            sources = [s for s in sources if _is_domain_allowed(s.get("url", ""))]
        response_payload = {
            "response": final_response_content,
            "activities": activities,
            "cost": {"tokens": len(final_response_content.split()), "estimated_cost": 0.002},
            "latency": 1.2,
            "triage": triage,
            "sources": sources,
            "classification": question_category,
            "triage_display": triage_display,
        }
        # Store in cache
        _cache_set(cache_key, response_payload)
        return response_payload
        
    except Exception as e:
        return {"error": f"Processing error: {str(e)}"}

# SSE endpoint for streaming-compatible environments (e.g., Vercel)
@app.get("/chat/stream")
async def chat_stream(message: str = "", session_id: Optional[str] = None, request: Request = None):
    """Server-Sent Events stream for chat. Keeps existing /chat unchanged.

    Query params:
      - message: user question
      - session_id: optional lab upload session
    """

    async def event_generator():
        # Activity: searching
        yield "event: activity\n" + "data: " + json.dumps({"tool": "BM25_Retriever", "status": "searching"}) + "\n\n"
        await asyncio.sleep(0.5)
        # Activity: processing
        yield "event: activity\n" + "data: " + json.dumps({"tool": "Medical_Agent", "status": "processing"}) + "\n\n"

        try:
            # Build messages similar to /chat
            messages = [HumanMessage(content=message)]
            lab_context = ""
            if session_id and session_id in uploaded_lab_data:
                lab_data = uploaded_lab_data[session_id]
                for file_data in lab_data["files"]:
                    content_preview = file_data["content"][:2000] if len(file_data["content"]) > 2000 else file_data["content"]
                    lab_context += f"File: {file_data['filename']}\nContent: {content_preview}\n\n"
                if lab_context.strip():
                    lab_instruction = (
                        "IMPORTANT: The user has uploaded lab results. Please analyze the uploaded lab data and provide a comprehensive medical interpretation.\n\n"
                        f"LAB RESULTS DATA:\n{lab_context}\n\nUSER QUESTION: {message}\n"
                    )
                    messages.insert(0, HumanMessage(content=lab_instruction))

            system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags.")
            final_response = compiled_health_graph.invoke({"messages": [system_message] + messages})
            final_message = final_response["messages"][-1]
            cleaned = (
                final_message.content.replace("**", "").replace("*", "").replace("###", "").replace("##", "").replace("#", "")
            )
            yield "event: response\n" + "data: " + json.dumps({"content": cleaned}) + "\n\n"
        except Exception as e:
            yield "event: error\n" + "data: " + json.dumps({"error": str(e)}) + "\n\n"

        # Heartbeat to keep connections healthy (one final ping)
        yield "event: ping\n" + "data: {}\n\n"

    headers = {
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "X-Accel-Buffering": "no",  # disable proxy buffering
    }
    return StreamingResponse(event_generator(), media_type="text/event-stream", headers=headers)

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    """WebSocket endpoint for real-time communication"""
    await manager.connect(websocket)
    try:
        while True:
            data = await websocket.receive_text()
            message = json.loads(data)
            
            # Process the medical query with real agent
            query = message.get("message", "")
            
            if not compiled_health_graph:
                await manager.send_personal_message(
                    json.dumps({"type": "error", "content": "Medical system not initialized"}), websocket
                )
                continue
            
            # Send real-time updates
            await manager.send_personal_message(
                json.dumps({
                    "type": "activity",
                    "tool": "BM25_Retriever",
                    "status": "searching",
                    "details": {"query": query}
                }), websocket
            )
            
            # Simulate processing time
            await asyncio.sleep(1)
            
            await manager.send_personal_message(
                json.dumps({
                    "type": "activity", 
                    "tool": "Medical_Agent",
                    "status": "processing",
                    "details": {"model": "gpt-4o-mini"}
                }), websocket
            )
            
            await asyncio.sleep(1)
            
            # Process with real agent
            try:
                system_message = SystemMessage(content="You are a helpful medical assistant. When providing links or references, use markdown format: [Link Text](URL). Do NOT use HTML tags like <a href=\"...\">. Keep responses clean and readable.")
                test_input = {"messages": [system_message, HumanMessage(content=query)]}
                final_response = compiled_health_graph.invoke(test_input)
                final_message = final_response["messages"][-1]
                
                # Send final response
                await manager.send_personal_message(
                    json.dumps({
                        "type": "response",
                        "content": final_message.content,
                        "cost": {"tokens": len(final_message.content.split()), "estimated_cost": 0.002},
                        "latency": 2.1
                    }), websocket
                )
                
            except Exception as e:
                await manager.send_personal_message(
                    json.dumps({
                        "type": "error",
                        "content": f"Processing error: {str(e)}"
                    }), websocket
                )
            
    except WebSocketDisconnect:
        manager.disconnect(websocket)

@app.get("/metrics")
async def get_metrics():
    """Get system metrics and performance data"""
    return {
        "status": "healthy",
        "active_connections": len(manager.active_connections),
        "total_requests": 0,  # Could be implemented with counters
        "system_uptime": "running"
    }

# -----------------------------
# Additive Enterprise Endpoints
# -----------------------------

@app.post("/v1/patient/summary", response_model=PatientSummaryResponse)
async def patient_summary(payload: PatientSummaryRequest):
    """Generate a structured patient summary from conversation history and optional lab session.
    Returns problems, medications, labs, and plan as JSON. Falls back to empty fields on errors.
    """
    try:
        messages = []
        for msg in payload.conversation_history:
            if msg.sender == "user":
                messages.append(HumanMessage(content=msg.text))
            else:
                messages.append(AIMessage(content=msg.text))

        # Attach lab context if present
        lab_context = ""
        if payload.session_id and payload.session_id in uploaded_lab_data:
            lab_data = uploaded_lab_data[payload.session_id]
            for file_data in lab_data["files"]:
                content_preview = file_data["content"][:1000]
                lab_context += f"File: {file_data['filename']}\nContent: {content_preview}\n\n"

        json_instruction = (
            "Return a concise JSON object with keys: problems (array of strings), medications (array of strings), "
            "labs (object with 'flags' array and 'notes' array), and plan (array of strings). "
            "Use ONLY the information provided in the conversation context. Do NOT invent symptoms or diagnoses. "
            "ALWAYS include a non-empty plan with 3-5 conservative, general actions that are appropriate to the problems/labs."
        )

        summary_prompt = f"""
        You are a clinical assistant. Analyze the conversation and produce a structured clinical summary.
        {('LAB CONTEXT:\n' + lab_context) if lab_context else ''}
        {json_instruction}
        Respond with valid JSON only.
        Conversation:\n{os.linesep.join([('User: ' + m.content) if isinstance(m, HumanMessage) else ('Assistant: ' + m.content) for m in messages])}
        """

        system_message = SystemMessage(content="You return only JSON. No prose. Keys: problems, medications, labs, plan.")
        llm_resp = compiled_health_graph.invoke({"messages": [system_message, HumanMessage(content=summary_prompt)]})
        content = llm_resp["messages"][-1].content

        try:
            parsed = json.loads(content)
        except Exception:
            # Attempt to extract JSON substring
            start = content.find("{")
            end = content.rfind("}")
            parsed = json.loads(content[start:end+1]) if start != -1 and end != -1 else {}

        problems = parsed.get("problems", []) or []
        medications = parsed.get("medications", []) or []
        labs_obj = parsed.get("labs", {}) or {}
        plan_list = parsed.get("plan", []) or []

        # If plan is empty, derive a conservative plan ONLY from the provided context
        if not plan_list:
            try:
                # Build a text blob from conversation history (frontend sends Snapshot lines)
                blob = "\n".join([m.text for m in payload.conversation_history or []]).lower()
                plan_list = ["Continue current medications as prescribed."]
                if "diabetes" in blob:
                    plan_list.append("Monitor blood glucose regularly; follow diet and exercise advice.")
                if "hypertension" in blob or "blood pressure" in blob:
                    plan_list.append("Monitor blood pressure; reduce sodium; follow lifestyle guidance.")
                if "asthma" in blob:
                    plan_list.append("Use rescue inhaler as directed; review triggers and inhaler technique.")
                if "last labs" in blob or "a1c" in blob or "ldl" in blob or "crp" in blob or "spirometry" in blob:
                    plan_list.append("Recheck relevant labs at the recommended interval.")
                plan_list.append("Schedule follow-up to reassess.")
            except Exception:
                plan_list = ["Continue current medications as prescribed.", "Schedule follow-up to reassess."]

        return PatientSummaryResponse(
            problems=problems,
            medications=medications,
            labs=labs_obj,
            plan=plan_list,
        )
    except Exception:
        return PatientSummaryResponse()


@app.post("/v1/labs/interpret", response_model=LabsInterpretResponse)
async def labs_interpret(payload: LabsInterpretRequest):
    """Interpret uploaded lab results for the given session and optional question.
    Uses existing in-memory uploaded_lab_data. Returns structured JSON.
    """
    try:
        if payload.session_id not in uploaded_lab_data:
            return LabsInterpretResponse(flags=["no_session"], explanations=["No lab session found"], recommendations=[])

        lab_data = uploaded_lab_data[payload.session_id]
        lab_text = "\n\n".join([f"{f['filename']}:\n{f['content'][:1500]}" for f in lab_data["files"]])
        q = payload.question or "Interpret the lab results for clinical significance."
        instruction = (
            "Return JSON with keys: flags (array of strings for abnormal findings), explanations (array of strings), "
            "recommendations (array of strings). Respond with JSON only."
        )

        system_message = SystemMessage(content="You are a medical assistant. Output strictly JSON as instructed.")
        llm_resp = compiled_health_graph.invoke({
            "messages": [
                system_message,
                HumanMessage(content=f"{instruction}\n\nLABS:\n{lab_text}\n\nQUESTION:\n{q}")
            ]
        })
        content = llm_resp["messages"][-1].content
        try:
            parsed = json.loads(content)
        except Exception:
            start = content.find("{")
            end = content.rfind("}")
            parsed = json.loads(content[start:end+1]) if start != -1 and end != -1 else {}

        return LabsInterpretResponse(
            flags=parsed.get("flags", []),
            explanations=parsed.get("explanations", []),
            recommendations=parsed.get("recommendations", []),
        )
    except Exception:
        return LabsInterpretResponse()


@app.post("/v1/provider/note", response_model=ProviderNoteResponse)
async def provider_note(payload: ProviderNoteRequest):
    """Generate a provider note (HPI/Assessment/Plan) from conversation messages.
    Returns both a sectioned object and a markdown string. Strictly additive to existing API.
    """
    try:
        sections = payload.sections or ["HPI", "Assessment", "Plan"]
        convo_text = "\n".join([f"{m.sender}: {m.text}" for m in payload.messages])
        snapshot_text = "" if not payload.snapshot else (
            f"Problems: {', '.join(payload.snapshot.get('problems', []))}\n"
            f"Medications: {', '.join(payload.snapshot.get('medications', []))}\n"
            f"Last Labs: {', '.join([f"{k}: {v}" for k, v in (payload.snapshot.get('last_labs', {}) or {}).items()])}"
        )
        visit_notes = payload.notes or ""
        instruction = (
            "Return a JSON object with keys matching sections (e.g., HPI, Assessment, Plan). "
            "Use ONLY the provided snapshot and visit notes. Do NOT invent symptoms or improvements. "
            "If a field is not supported by the notes/snapshot, keep it concise and factual. Respond with JSON only."
        )
        system_message = SystemMessage(content="You are a clinical note assistant. Output JSON only. Be strictly factual to provided inputs.")
        llm_resp = compiled_health_graph.invoke({
            "messages": [
                system_message,
                HumanMessage(content=f"{instruction}\n\nSECTIONS: {', '.join(sections)}\n\nSNAPSHOT:\n{snapshot_text}\n\nVISIT_NOTES:\n{visit_notes}\n\nCONVERSATION:\n{convo_text}")
            ]
        })
        content = llm_resp["messages"][-1].content
        try:
            note_obj = json.loads(content)
        except Exception:
            start = content.find("{")
            end = content.rfind("}")
            note_obj = json.loads(content[start:end+1]) if start != -1 and end != -1 else {}

        md_lines = []
        for sec in sections:
            value = note_obj.get(sec, "")
            md_lines.append(f"## {sec}\n{value}\n")

        return ProviderNoteResponse(note=note_obj, markdown="\n".join(md_lines))
    except Exception as e:
        return ProviderNoteResponse(note={}, markdown=f"Error: {str(e)}")


@app.get("/v1/demo/patients")
async def demo_patients():
    """Demo patients list for provider UI without persistence."""
    return [
        {"id": "p1", "name": "Alex Johnson", "dob": "1985-03-14", "mrn": "AJ-0001"},
        {"id": "p2", "name": "Maria Gomez", "dob": "1972-09-02", "mrn": "MG-0042"},
        {"id": "p3", "name": "Sam Lee", "dob": "1991-12-20", "mrn": "SL-0099"},
    ]


@app.get("/v1/demo/patients/{patient_id}")
async def demo_patient_detail(patient_id: str):
    """Demo patient detail for provider UI."""
    demo = {
        "p1": {
            "demographics": {"age": 39, "sex": "M"},
            "problems": ["Type 2 Diabetes", "Hypertension"],
            "medications": ["Metformin 500mg BID", "Lisinopril 10mg QD"],
            "last_labs": {"A1C": "7.8%", "LDL": "130 mg/dL"},
        },
        "p2": {
            "demographics": {"age": 52, "sex": "F"},
            "problems": ["Rheumatoid Arthritis"],
            "medications": ["Methotrexate 15mg weekly"],
            "last_labs": {"CRP": "12 mg/L"},
        },
        "p3": {
            "demographics": {"age": 33, "sex": "M"},
            "problems": ["Asthma"],
            "medications": ["Albuterol PRN"],
            "last_labs": {"Spirometry": "Mild obstruction"},
        },
    }
    return demo.get(patient_id, {"error": "not_found"})

@app.post("/generate-title")
async def generate_title(request: GenerateTitleRequest):
    """Generate intelligent conversation title using LLM"""
    try:
        # Create a prompt for title generation
        messages_text = "\n".join([f"{msg['sender']}: {msg['text']}" for msg in request.messages])
        
        title_prompt = f"""
        Analyze this medical conversation and generate a concise, meaningful title (maximum 5 words).
        Focus on the specific medical condition, symptom, or health topic discussed.
        
        Examples:
        - "tell me about diabetes" → "Diabetes Information"
        - "what are symptoms of heart disease" → "Heart Disease Symptoms"
        - "how to manage blood pressure" → "Blood Pressure Management"
        
        Conversation:
        {messages_text}
        
        Generate only the title, nothing else:
        """
        
        # Use LLM to generate title
        system_message = SystemMessage(content="You are a medical assistant. Generate concise, accurate titles for medical conversations. Focus on the specific medical condition or health topic discussed. Do not use generic terms like 'Health Consultation' or 'Pet Health' unless the conversation is actually about pets.")
        response = compiled_health_graph.invoke({
            "messages": [system_message, HumanMessage(content=title_prompt)]
        })
        
        title = response["messages"][-1].content.strip()
        
        # Clean up the title
        title = title.replace('"', '').replace("'", "").strip()
        if title.startswith("Title:"):
            title = title[6:].strip()
        
        # Fallback if LLM fails
        if not title or len(title) > 50:
            title = "Health Consultation"
            
        return {"title": title}
        
    except Exception as e:
        return {"title": "Health Consultation"}

@app.post("/upload-lab-results")
async def upload_lab_results(files: List[UploadFile] = File(...), request: Request = None):
    """Upload and process lab result files"""
    try:
        # Rate limit per-IP for uploads
        client_ip = request.client.host if request and request.client else "unknown"
        if not check_rate_limit(f"upload:{client_ip}", UPLOAD_LIMIT_PER_MINUTE):
            raise HTTPException(status_code=429, detail="Upload rate limit exceeded. Please wait and try again.")

        session_id = str(uuid.uuid4())
        processed_files = []
        
        for file in files:
            if file.filename:
                text_content = process_uploaded_file(file)
                print(f"DEBUG: Extracted content from {file.filename}: {text_content[:200]}...")  # Debug line
                processed_files.append({
                    "filename": file.filename,
                    "content": text_content
                })
        
        # Store processed data
        uploaded_lab_data[session_id] = {
            "files": processed_files,
            "timestamp": asyncio.get_event_loop().time()
        }
        
        return UploadResponse(
            session_id=session_id,
            files_processed=len(processed_files),
            message=f"Successfully processed {len(processed_files)} file(s)"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.get("/lab-results/{session_id}")
async def get_lab_results(session_id: str):
    """Retrieve processed lab results for a session"""
    if session_id not in uploaded_lab_data:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return uploaded_lab_data[session_id]

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 